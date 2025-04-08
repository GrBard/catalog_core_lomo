import os
import pandas as pd
import re
from pathlib import Path
from app.utils import find_continuous_intervals, resource_path
from docx import Document
from docx.shared import Inches, Cm
from PIL import Image, ImageDraw, ImageFont
import io
import math
import time

class DataProcessor:
    def __init__(self, excel_path, images_folder, box_column="BOX", start_column="от", end_column="до", measurements_column="замеры"):
        print(f"Инициализация DataProcessor: box_column='{box_column}', start_column='{start_column}', end_column='{end_column}', measurements_column='{measurements_column}'")
        self.excel_path = Path(excel_path).resolve()
        self.images_folder = Path(images_folder).resolve()
        self.data = None
        self.all_image_files = []
        self.current_dataframe = None
        self.box_column = box_column
        self.start_column = start_column
        self.end_column = end_column
        self.measurements_column = measurements_column

    def load_excel(self):
        """Загружает данные из Excel в DataFrame."""
        self.data = pd.read_excel(self.excel_path)

    def load_image_files(self):
        """Получает список всех изображений из выбранной папки."""
        valid_extensions = (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp")
        self.all_image_files = [f.resolve() for f in self.images_folder.iterdir() if f.suffix.lower() in valid_extensions]

    def add_photo_columns(self):
        """Добавляет столбцы с путями к фотографиям и названием скважины."""
        if not self.all_image_files:
            self.load_image_files()

        if self.box_column not in self.data.columns:  # Используем self.box_column вместо 'BOX'
            raise ValueError(f"В данных отсутствует столбец '{self.box_column}' для сопоставления с фото.")

        def find_matching_file(box, uf=False):
            box_str = str(box)
            pattern = re.compile(rf"^(.+?)_{box_str}(_uf)?\.(jpg|jpeg|png|tif|tiff|bmp)$", re.IGNORECASE)
            for file_path in self.all_image_files:
                file_name = file_path.name.lower()
                match = pattern.match(file_name)
                if match:
                    if uf and "_uf" in file_name:
                        return str(file_path)
                    elif not uf and "_uf" not in file_name:
                        return str(file_path)
            return None

        def extract_well_name(photo_path):
            if not photo_path:
                return None
            file_name = Path(photo_path).stem
            match = re.match(r"^(.+?)_\d+(_uf)?$", file_name)
            if match:
                well_name = match.group(1)
                return well_name.replace("скв.", "").strip()
            return None

        self.data["Фото"] = self.data[self.box_column].apply(lambda box: find_matching_file(box, uf=False))
        self.data["Фото УФ"] = self.data[self.box_column].apply(lambda box: find_matching_file(box, uf=True))
        self.data["Скважина"] = self.data["Фото"].apply(extract_well_name)
        cols = list(self.data.columns)
        cols.insert(cols.index(self.box_column), cols.pop(cols.index("Скважина")))
        self.data = self.data[cols]

    def compute_intervals(self):
        """Вычисляет непрерывные интервалы и добавляет их в DataFrame как 'Начало интервала' и 'Конец интервала'."""
        cols_lower = {col.lower(): col for col in self.data.columns}
        if self.start_column.lower() not in cols_lower or self.end_column.lower() not in cols_lower:
            raise ValueError(
                f"В данных отсутствуют столбцы '{self.start_column}' или '{self.end_column}' для вычисления интервалов.")

        start_col = cols_lower[self.start_column.lower()]
        end_col = cols_lower[self.end_column.lower()]

        _, interval_dict = find_continuous_intervals(self.data, start_col, end_col)

        def get_interval_start(row):
            start = float(row[start_col])
            return interval_dict.get(start, (None, None))[0]

        def get_interval_end(row):
            start = float(row[start_col])
            return interval_dict.get(start, (None, None))[1]

        self.data["Начало интервала"] = self.data.apply(get_interval_start, axis=1)
        self.data["Конец интервала"] = self.data.apply(get_interval_end, axis=1)

        cols = list(self.data.columns)
        end_col_index = cols.index(end_col)
        cols.pop(cols.index("Начало интервала"))
        cols.pop(cols.index("Конец интервала"))
        cols.insert(end_col_index + 1, "Начало интервала")
        cols.insert(end_col_index + 2, "Конец интервала")
        self.data = self.data[cols]

    def process_data(self, samples_file=None):
        """Обрабатывает данные: загружает Excel, добавляет фото, вычисляет интервалы."""
        self.load_excel()
        self.add_photo_columns()
        self.compute_intervals()

        # Вычисляем "Вынос"
        cols_lower = {col.lower(): col for col in self.data.columns}
        required_cols_lower = [self.measurements_column.lower(), self.end_column.lower(),
                               self.start_column.lower()]  # Обновлено
        missing_cols = [col for col in required_cols_lower if col not in cols_lower]

        if not missing_cols:
            measurements_col = cols_lower[self.measurements_column.lower()]  # Используем динамическое имя
            end_col = cols_lower[self.end_column.lower()]
            start_col = cols_lower[self.start_column.lower()]

            self.data[measurements_col] = pd.to_numeric(self.data[measurements_col], errors="coerce")
            self.data[end_col] = pd.to_numeric(self.data[end_col], errors="coerce")
            self.data[start_col] = pd.to_numeric(self.data[start_col], errors="coerce")

            self.data["Вынос"] = self.data.apply(
                lambda
                    row: f"{row[measurements_col]} м ({(row[measurements_col] / (row[end_col] - row[start_col]) * 100):.1f} %)"
                if pd.notna(row[measurements_col]) and pd.notna(row[end_col]) and pd.notna(row[start_col]) and (
                        row[end_col] - row[start_col]) != 0
                else "N/A",
                axis=1
            )
            cols = list(self.data.columns)
            cols.insert(cols.index(measurements_col) + 1, cols.pop(cols.index("Вынос")))  # Обновлено
            self.data = self.data[cols]

        self.current_dataframe = self.data.copy()
        if self.start_column in self.current_dataframe.columns:
            self.current_dataframe = self.current_dataframe.sort_values(by=self.start_column)
        return self.current_dataframe

    def generate_depth_scale(self, top_depth, bottom_depth, core_count, box_length=1.0):
        """Генерирует шкалу глубин для коробки."""
        # Параметры шкалы
        shift_up = 20  # Сдвиг сверху в пикселях
        shift_btm = 20  # Сдвиг снизу в пикселях
        image_height = 1100  # Высота изображения шкалы (пиксели)
        image_width = 50  # Ширина изображения шкалы
        step_rul = 0.1  # Шаг линейки (0.1 м)

        # Создаём изображение для первой шкалы
        im_d = Image.new('RGB', (image_width, image_height), (255, 255, 255))
        draw = ImageDraw.Draw(im_d)

        # Если есть вторая шкала (core_count == 2), создаём её
        if core_count == 2:
            im_d2 = Image.new('RGB', (image_width, image_height), (255, 255, 255))
            draw2 = ImageDraw.Draw(im_d2)
        else:
            draw2 = None

        # Загружаем шрифт
        try:
            font_path = resource_path('resources/arial.ttf')
            font = ImageFont.truetype(font_path, 14)
            font_small = ImageFont.truetype(font_path, 10)
        except IOError:
            font = ImageFont.load_default()
            font_small = ImageFont.load_default()

        # Определяем направление шкалы (сверху вниз)
        rul_direct = 1  # 1: сверху вниз, -1: снизу вверх
        if rul_direct == 1:
            shift = shift_up
            h = (image_height - shift_up - shift_btm) / (box_length * 10)  # Масштаб в пикселях на дециметр
            dy_text = 15  # Смещение текста для первой отметки
        else:
            shift = image_height - shift_btm
            h = -(image_height - shift_up - shift_btm) / (box_length * 10)
            dy_text = -15

        # Рисуем отметки на шкале
        for depth in [top_depth, bottom_depth]:
            for ll in range(round((bottom_depth - top_depth) / step_rul) + 1):
                current_depth = top_depth + ll * step_rul
                # Вычисляем позицию на шкале
                dz = (current_depth - top_depth) / box_length  # Нормализованная позиция
                hy = (image_height - shift_up - shift_btm) * dz
                tsh = dy_text if ll == 0 else 0  # Смещение текста для первой отметки

                # Определяем, на какую шкалу рисовать (первая или вторая)
                target_draw = draw if ll < (round((bottom_depth - top_depth) / step_rul) + 1) // 2 or core_count == 1 else draw2

                # Рисуем отметки
                if current_depth % 1 < 0.00001:  # Каждые 1 м
                    target_draw.line((0, shift_up + hy, 50, shift_up + hy), fill="green", width=4)
                    target_draw.text((0, shift_up + hy + tsh), str(round(current_depth, 1)), (0, 0, 0), font=font)
                elif current_depth % 0.5 < 0.00001:  # Каждые 0.5 м
                    target_draw.line((0, shift_up + hy, 50, shift_up + hy), fill="green", width=2)
                    target_draw.text((0, shift_up + hy + tsh), str(round(current_depth, 1)), (0, 0, 0), font=font)
                else:  # Каждые 0.1 м
                    target_draw.line((0, shift_up + hy, 25, shift_up + hy), fill="green", width=2)

        # Сохраняем первую шкалу в поток
        img_d = io.BytesIO()
        im_d.save(img_d, 'JPEG')
        img_d.seek(0)

        # Если есть вторая шкала, сохраняем её
        if core_count == 2:
            img_d2 = io.BytesIO()
            im_d2.save(img_d2, 'JPEG')
            img_d2.seek(0)
            return img_d, img_d2
        return img_d, None

    def draw_sample_circles(self, photo_path, samples_in_box, core_count, suffix='_with_circles'):
        """Рисует кружки на копии фото керна в местах отбора образцов, не изменяя оригинал."""
        # Открываем изображение
        img = Image.open(photo_path)
        # Создаём копию изображения
        img_copy = Image.new('RGB', img.size)
        img_copy.paste(img)  # Копируем содержимое, чтобы гарантированно не изменять оригинал
        draw = ImageDraw.Draw(img_copy)

        # Параметры изображения
        img_width, img_height = img_copy.size
        shift_up = 0  # Отступ сверху (можно настроить, если нужно)
        shift_btm = 0  # Отступ снизу (можно настроить, если нужно)

        # Радиус кружка (12% от ширины изображения)
        r = img_width * 0.12

        # Вычисляем позицию по горизонтали (центр для одного ядра)
        hx = img_width / (2 * core_count)

        # Настраиваем шрифт для подписи
        try:
            font_path = resource_path('resources/arial.ttf')
            font = ImageFont.truetype(font_path, int(r * 1))
        except IOError:
            font = ImageFont.load_default()

        # Отступы для текста
        sht_txt_g = r  # Горизонтальный отступ текста
        sht_txt_v = r  # Вертикальный отступ текста

        # Обрабатываем каждый образец в коробке
        for _, sample in samples_in_box.iterrows():
            sample_num = sample['Номер образца']
            # Извлекаем десятичную часть номера образца (глубина в метрах от верха коробки)
            depth_in_box = sample_num - int(sample_num)  # Например, для 3.15 это 0.15 м

            # Вычисляем вертикальную позицию кружка
            # Коробка 1 м, высота изображения (за вычетом отступов) соответствует 1 м
            hy = (img_height - shift_up - shift_btm) * depth_in_box

            # Корректируем положение, если кружок близко к краю
            if shift_up + hy + r > 0.95 * img_height:  # Близко к нижнему краю
                hy = 0.95 * img_height - shift_up - r
            elif shift_up + hy - r < 0.05 * img_height:  # Близко к верхнему краю
                hy = 0.05 * img_height - shift_up + r

            # Рисуем кружок
            draw.arc(
                (hx - r, shift_up + hy - r, hx + r, shift_up + hy + r),
                0, 360, fill=(255, 255, 0), width=int(r / 10)
            )

            # Добавляем подпись (номер образца)
            if shift_up + hy + r > 0.95 * img_height:
                # Если кружок внизу, подпись выше
                draw.text(
                    (hx - sht_txt_g, shift_up + hy - r - sht_txt_v),
                    str(sample_num),
                    (255, 255, 0),
                    font=font
                )
            else:
                # Иначе подпись ниже
                draw.text(
                    (hx - sht_txt_g, shift_up + hy + r),
                    str(sample_num),
                    (255, 255, 0),
                    font=font
                )

        # Закрываем исходное изображение
        img.close()

        # Сохраняем изменённое изображение во временный файл
        temp_img = io.BytesIO()
        img_copy.save(temp_img, 'PNG')
        temp_img.seek(0)

        # Создаём уникальный временный путь, чтобы избежать перезаписи
        base, ext = os.path.splitext(photo_path)
        temp_path = f"{base}{suffix}_{int(time.time())}{ext}"  # Добавляем временную метку для уникальности
        with open(temp_path, 'wb') as f:
            f.write(temp_img.getvalue())

        # Проверяем, что исходный файл не изменился
        if os.path.getsize(photo_path) != os.path.getsize(temp_path):
            print(f"Исходный файл {photo_path} не изменён.")
        else:
            print(f"Внимание: исходный файл {photo_path} может быть изменён!")

        return temp_path

    def create_catalog(self, save_path, samples_df=None, progress_bar=None, progress_step=1.0):
        print("Внутри DataProcessor.create_catalog")
        print(f"Используемый box_column: '{self.box_column}'")
        if self.current_dataframe is None:
            raise ValueError("Нет данных для создания каталога.")

        print(f"Столбцы в current_dataframe: {list(self.current_dataframe.columns)}")
        if self.box_column not in self.current_dataframe.columns:
            raise ValueError(f"Столбец '{self.box_column}' отсутствует в данных.")

        scale_image_path = resource_path('resources/scale.jpg')
        shkala_image_path = resource_path('resources/shkala.jpg')

        if not os.path.exists(scale_image_path):
            raise FileNotFoundError(f"Файл {scale_image_path} не найден.")
        if not os.path.exists(shkala_image_path):
            raise FileNotFoundError(f"Файл {shkala_image_path} не найден.")

        doc = Document()
        print("Документ создан")

        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        doc.add_heading(f'Фотографии керна по скважине {self.current_dataframe["Скважина"].iloc[0]}', 0)
        print("Заголовок добавлен")

        p = doc.add_paragraph('Глубины даны по керну. ')
        p.add_run('Номера образцов расположены напротив точек выбуривания. ')
        p.add_run('Номер образца состоит из двух цифр, разделённых точкой. ')
        p.add_run(
            'Первая цифра соответствует номеру коробки, вторая – расстояние в сантиметрах от низа коробки до точки отбора образца.')
        print("Вступительный текст добавлен")

        width_samles_col = 0.9
        width_photo_col = 3.5
        width_samles_right = 4.0
        target_height = Inches(8.614)
        shkala_height = Inches(1)

        grouped = self.current_dataframe.groupby(self.box_column)
        print(f"Группировка выполнена, групп: {len(grouped)}")

        cols_lower = {col.lower(): col for col in self.current_dataframe.columns}
        start_col = cols_lower.get(self.start_column.lower())
        end_col = cols_lower.get(self.end_column.lower())

        if not start_col or not end_col:
            raise ValueError(f"Не найдены столбцы '{self.start_column}' и/или '{self.end_column}' в DataFrame.")

        for box_number, group in grouped:
            print(f"Обработка коробки {box_number}")
            doc.add_page_break()

            table = doc.add_table(rows=4, cols=4, style='Table Grid')
            for cell in table.columns[0].cells:
                cell.width = Inches(width_samles_col)
            for cell in table.columns[1].cells:
                cell.width = Inches(width_photo_col)
            for cell in table.columns[2].cells:
                cell.width = Inches(width_samles_right)
            for cell in table.columns[3].cells:
                cell.width = Inches(0.5)

            cell = table.cell(0, 0)
            cell.text = f'Коробка {int(box_number)}'

            cell = table.cell(0, 1)
            ts = 'Интервал бурения: '
            for idx, row in group.iterrows():
                ts += f"{row['Начало интервала']}-{row['Конец интервала']}\nвынос: {row['Вынос']}\n"
            cell.text = ts.strip()

            cell = table.cell(1, 0)
            cell.text = 'Номера образцов:'

            cell = table.cell(2, 0)
            if samples_df is not None:
                samples_in_box = samples_df[samples_df[self.box_column] == int(box_number)]
                sample_numbers = samples_in_box['Номер образца'].tolist()
                cell.text = '\n'.join(map(str, sample_numbers))
            else:
                cell.text = ''

            cell = table.cell(1, 1)
            start_value = group[start_col].iloc[0]
            cell.text = f'[{start_value}]'

            cell = table.cell(3, 1)
            end_value = group[end_col].iloc[0]
            cell.text = f'[{end_value}]'

            cell = table.cell(1, 2)
            cell.text = 'Исследования:'

            cell = table.cell(2, 2)
            if samples_df is not None and not samples_in_box.empty:
                samples_in_box = samples_in_box.sort_values(by='Номер образца')
                for _, sample in samples_in_box.iterrows():
                    paragraph = cell.add_paragraph()
                    run = paragraph.add_run()
                    run.text = f"{sample['Номер образца']}"
                    run.add_tab()
                    run.text += f"{sample['Исследования']}"

            # Добавляем шкалу глубин, основное фото и УФ-фото
            cell = table.cell(2, 1)
            paragraph = cell.paragraphs[0]
            top_depth = float(group[start_col].iloc[0])
            bottom_depth = float(group[end_col].iloc[0])
            core_count = 1
            print(f"Генерация шкалы глубин для коробки {box_number}: {top_depth} - {bottom_depth}")
            img_d, img_d2 = self.generate_depth_scale(top_depth, bottom_depth, core_count)
            run = paragraph.add_run()
            run.add_picture(img_d, height=target_height)
            if img_d2:
                run = paragraph.add_run()
                run.add_picture(img_d2, height=target_height)

            photo_path = group["Фото"].iloc[0]
            if pd.notna(photo_path) and os.path.exists(photo_path):
                if samples_df is not None and not samples_in_box.empty:
                    print(f"Обрабатываем основное фото для коробки {box_number}: {photo_path}")
                    modified_photo_path = self.draw_sample_circles(photo_path, samples_in_box, core_count)
                    print(f"Добавляем основное фото с кружками для коробки {box_number}: {modified_photo_path}")
                    run = paragraph.add_run()
                    run.add_picture(modified_photo_path, height=target_height)
                    try:
                        os.remove(modified_photo_path)
                        print(f"Удалён временный файл основного фото: {modified_photo_path}")
                    except Exception as e:
                        print(f"Ошибка удаления временного файла основного фото: {e}")
                else:
                    print(f"Добавляем оригинальное основное фото для коробки {box_number}: {photo_path}")
                    run = paragraph.add_run()
                    run.add_picture(photo_path, height=target_height)
            else:
                print(f"Основное фото для коробки {box_number} не найдено или отсутствует: {photo_path}")

            print(f"Добавляем шкалу для коробки {box_number}: {shkala_image_path}")
            run = paragraph.add_run()
            run.add_picture(shkala_image_path, height=shkala_height)

            photo_uf_path = group["Фото УФ"].iloc[0]
            if pd.notna(photo_uf_path) and os.path.exists(photo_uf_path):
                if samples_df is not None and not samples_in_box.empty:
                    print(f"Обрабатываем УФ-фото для коробки {box_number}: {photo_uf_path}")
                    modified_uf_photo_path = self.draw_sample_circles(photo_uf_path, samples_in_box, core_count,
                                                                      suffix='_uf_with_circles')
                    print(f"Добавляем УФ-фото с кружками для коробки {box_number}: {modified_uf_photo_path}")
                    run = paragraph.add_run()
                    run.add_picture(modified_uf_photo_path, height=target_height)
                    try:
                        os.remove(modified_uf_photo_path)
                        print(f"Удалён временный файл УФ-фото: {modified_uf_photo_path}")
                    except Exception as e:
                        print(f"Ошибка удаления временного файла УФ-фото: {e}")
                else:
                    print(f"Добавляем оригинальное УФ-фото для коробки {box_number}: {photo_uf_path}")
                    run = paragraph.add_run()
                    run.add_picture(photo_uf_path, height=target_height)
            else:
                print(f"УФ-фото для коробки {box_number} не найдено или отсутствует: {photo_uf_path}")

            # Добавляем масштаб
            cell = table.cell(2, 3)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            print(f"Добавляем масштаб для коробки {box_number}: {scale_image_path}")
            run.add_picture(scale_image_path, height=target_height)

            if progress_bar is not None:
                progress_bar.set(min(progress_bar.get() + progress_step, 1.0))
                progress_bar.update()

        print(f"Сохранение в {save_path}")
        doc.save(save_path)
        print(f"Документ сохранён: {save_path}")
        return save_path

    def get_current_dataframe(self):
        """Возвращает текущий DataFrame."""
        return self.current_dataframe
